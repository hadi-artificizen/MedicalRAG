import os
import re
import shutil
import time
import csv
from datetime import datetime
from typing import List
import psycopg2
from dotenv import load_dotenv
from langchain_community.document_loaders import PyPDFLoader, CSVLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
from langchain.schema import Document
from langchain_openai import ChatOpenAI
from langchain.schema import HumanMessage, SystemMessage
from docx import Document as DocxDocument
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import bcrypt
import pytesseract
from PIL import Image
import whisper
from moviepy.editor import VideoFileClip
from pydub import AudioSegment
from langchain.tools import tool, StructuredTool
from langchain.agents import initialize_agent, AgentType
from pydantic import BaseModel, Field
from typing import TypedDict, Optional
import json
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from io import BytesIO
import fitz
load_dotenv()

DB_URL = os.getenv("DB_URL")
conn = psycopg2.connect(DB_URL)
cur = conn.cursor()

splitter = RecursiveCharacterTextSplitter(chunk_size=300, chunk_overlap=100)

chat_llm = ChatOpenAI(
    model="openai/gpt-oss-120b",
    temperature=0.5,
    base_url="https://api.groq.com/openai/v1",
    api_key=os.getenv("OPENAI_API_KEY")
)

SMTP_USER = "imabdul.hadi1234@gmail.com"
SMTP_PASSWORD = "gqty jnji ufjq rgrh"
EMAIL_TIMEOUT = 30

PERSIST_DIR = "D:/LangchainPractice/RAG/chroma_db"

ASSETS_DIR = os.getenv("ASSETS_DIR", os.path.join(os.getcwd(), "assets"))
os.makedirs(ASSETS_DIR, exist_ok=True)

ASSET_SUBFOLDERS = {
    ".pdf": "pdfs",
    ".docx": "docs",
    ".csv": "csvs",
    ".png": "images",
    ".jpg": "images",
    ".jpeg": "images",
    ".gif": "images",
    ".bmp": "images",
    ".tiff": "images",
    ".mp4": "videos",
    ".mp3": "audios"
}

def sanitize_filename(filename: str) -> str:
    filename = os.path.basename(filename)
    filename = re.sub(r'[^A-Za-z0-9_.\-() ]+', '_', filename)
    return filename

def save_uploaded_file(tmp_path: str, original_filename: str) -> str:
    """
    Move/copy the temporary uploaded file into assets/<type> and return absolute path.
    The asset_uri will be built in load_docs_from_file for API serving.
    """
    ext = os.path.splitext(original_filename.lower())[1]
    sub = ASSET_SUBFOLDERS.get(ext, "others")
    dest_dir = os.path.join(ASSETS_DIR, sub)
    os.makedirs(dest_dir, exist_ok=True)
    safe_name = f"{int(time.time())}_{sanitize_filename(original_filename)}"
    dest_path = os.path.join(dest_dir, safe_name)
    shutil.copyfile(tmp_path, dest_path)
    # Return absolute filesystem path
    return dest_path
# def save_uploaded_file(tmp_path: str, original_filename: str) -> str:
#     """
#     Move/copy the temporary uploaded file into assets/<type> and return absolute path.
#     Metadata will contain 'asset_uri' set to a path under /assets/... (relative).
#     """
#     ext = os.path.splitext(original_filename.lower())[1]
#     sub = ASSET_SUBFOLDERS.get(ext, "others")
#     dest_dir = os.path.join(ASSETS_DIR, sub)
#     os.makedirs(dest_dir, exist_ok=True)
#     safe_name = f"{int(time.time())}_{sanitize_filename(original_filename)}"
#     dest_path = os.path.join(dest_dir, safe_name)
#     shutil.copyfile(tmp_path, dest_path)
#     # Return absolute filesystem path (load_docs_from_file will build asset_uri)
#     return dest_path

def hash_password(password: str) -> str:
    return bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')

def verify_password(plain_password: str, hashed_password: str) -> bool:
    return bcrypt.checkpw(plain_password.encode('utf-8'), hashed_password.encode('utf-8'))

def send_email(to_email , subject , body) :
    """Send an email to the given address with subject and body. Call this function for ANY medical query response."""
    print("Email function called")
    if not to_email or "@" not in to_email:
        return {"success": False, "message": "Invalid email address"}
    
    if not body.strip():
        return {"success": False, "message": "Email body cannot be empty"}
    
    msg = MIMEMultipart()
    msg["Subject"] = subject
    msg["From"] = SMTP_USER
    msg["To"] = to_email
    msg.attach(MIMEText(body, "plain"))
    
    try:
        with smtplib.SMTP_SSL("smtp.gmail.com", 465, timeout=EMAIL_TIMEOUT) as server:
            print("SMTP connection established")
            server.login(SMTP_USER, SMTP_PASSWORD)
            server.send_message(msg)
        print(f"Email sent successfully to {to_email}")
        return {"messages" : f"success : True, message : Email sent to {to_email}"}
    except Exception as e:
        error_msg = f"Email failed: {str(e)}"
        print(error_msg)
        return {"messages" : f"success : False, message : {error_msg}"}

class SendEmailInput(BaseModel):
    to_email: str = Field(default="", description="Recipient email address")
    subject: str = Field(default="Medical Query Response", description="Subject of the email")
    body: str = Field(..., description="Body of the email containing the medical response")

send_email_tool = StructuredTool.from_function(
    func=send_email,
    name="send_email",
    description="Send an email with the medical response. ALWAYS call this function after providing any medical advice or response.",
    args_schema=SendEmailInput
)


from langgraph.graph import StateGraph, END


llm_with_tools = chat_llm.bind_tools([send_email_tool])

class graphState(TypedDict):
    query : str
    doc : str
    mail : Optional[str]
    email_body : Optional[str]
    subject : Optional[str]

graph = StateGraph(graphState)

def agent_node(state):

    query = state["query"]
    doc = state["doc"]


graph.add_node("agent", agent_node)
graph.set_entry_point("agent")
graph.add_edge("agent", END)

# Compile graph
agent = graph.compile()

def get_embeddings():
    return HuggingFaceEmbeddings(
        model_name="abhinand/MedEmbed-base-v0.1",
        model_kwargs={"device": "cpu"}
    )

def get_vectorstore():
    return Chroma(embedding_function=get_embeddings(), persist_directory=PERSIST_DIR)


def _page_doc_metadata_for_pdf(doc: Document, meta_base: dict, idx: int):
    # PyPDFLoader often sets metadata['page']; unify to 'page'
    m = dict(meta_base)
    m["doc_type"] = "pdf"
    page_num = doc.metadata.get("page") if isinstance(doc.metadata, dict) else None
    if page_num is None:
        page_num = idx + 1
    m["page"] = int(page_num+1)
    m["full_text"] = doc.page_content
    return m

def load_docs_from_file(tmp_path: str, ext: str, source: str) -> List[Document]:
    

    meta_base = {"source": source, "processed_date": datetime.now().isoformat()}
    docs: List[Document] = []

   
    if os.path.commonpath([os.path.abspath(tmp_path), os.path.abspath(ASSETS_DIR)]) != os.path.abspath(ASSETS_DIR):
        # tmp_path not in assets dir -> copy it
        try:
            saved_path = save_uploaded_file(tmp_path, source)
        except Exception as e:
            print("Failed to save uploaded file to assets:", e)
            saved_path = tmp_path
    else:
        saved_path = tmp_path


    try:
        rel_path = os.path.relpath(saved_path, ASSETS_DIR).replace("\\", "/")
        asset_uri = f"assets/{rel_path}"  # Remove leading slash for API endpoint
    except Exception:
        asset_uri = f"assets/{os.path.basename(saved_path)}"

    meta_base["asset_uri"] = asset_uri
    # try:
    #     rel_path = os.path.relpath(saved_path, ASSETS_DIR).replace("\\", "/")
    #     asset_uri = f"/assets/{rel_path}"
    # except Exception:
    #     asset_uri = f"/assets/{os.path.basename(saved_path)}"

    # meta_base["asset_uri"] = asset_uri

    ext = ext.lower()

    if ext == ".pdf":
        try:
            pdf_docs = PyPDFLoader(saved_path).load()  # returns per-page Documents
            for idx, pd in enumerate(pdf_docs):
                m = _page_doc_metadata_for_pdf(pd, meta_base, idx)
                d = Document(page_content=pd.page_content, metadata=m)
                docs.append(d)
        except Exception as e:
            print(f"PDF load error for {source}: {e}")
            return []

    elif ext == ".docx":
        try:
            doc = DocxDocument(saved_path)
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]

            # Step 1: Create PDF from DOCX (like your render function)
            pdf_buffer = BytesIO()
            pdf_doc = SimpleDocTemplate(pdf_buffer, pagesize=letter)
            styles = getSampleStyleSheet()
            story = [Paragraph(p, styles['Normal']) for p in paragraphs]
            pdf_doc.build(story)
            pdf_buffer.seek(0)

            # Step 2: Load PDF with fitz to map paragraphs → pages
            pdf_document = fitz.open(stream=pdf_buffer.getvalue(), filetype="pdf")

            paragraph_page_map = {}
            for page_num in range(pdf_document.page_count):
                page_text = pdf_document[page_num].get_text("text")
                for idx, para in enumerate(paragraphs):
                    if para and para in page_text:
                        paragraph_page_map[idx] = page_num

            pdf_document.close()
            pdf_buffer.close()

            # Step 3: Store docs with paragraph_index + page_number
            for idx, para in enumerate(paragraphs):
                m = dict(meta_base)
                m["doc_type"] = "docx"
                m["paragraph_index"] = idx
                m["full_text"] = para
                m["docx_url"] = f"/assets/docs/{os.path.basename(saved_path)}"
                m["page_number"] = paragraph_page_map.get(idx, None)  # save page number if found

                docs.append(Document(page_content=para, metadata=m))

        except Exception as e:
            print(f"DOCX load error for {source}: {e}")
            return []
    
    elif ext == ".csv":
        try:
            with open(saved_path, "r", encoding="utf-8-sig") as fh:
                reader = csv.DictReader(fh)
                headers = reader.fieldnames or []
                for idx, row in enumerate(reader, start=1):
                    text = " | ".join(f"{k}: {v}" for k, v in row.items())

                    m = dict(meta_base)
                    m["doc_type"] = "csv"
                    m["row_number"] = idx
                    m["csv_columns"] = ",".join(headers)
                    m["snippet"] = text[:300]

                    # ✅ Save full row safely as JSON string
                    import json
                    m["row_data"] = json.dumps(row, ensure_ascii=False)

                    # Each row → separate Document
                    docs.append(Document(page_content=text, metadata=m))
        except Exception as e:
            print(f"CSV load error for {source}: {e}")
            return []
        except Exception as e:
            print(f"CSV load error for {source}: {e}")
            return []

    elif ext in [".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff"]:
        try:
            img = Image.open(saved_path)
            text = pytesseract.image_to_string(img)
            m = dict(meta_base)
            m["doc_type"] = "image"
            m["snippet"] = (text or "")[:300]
            m["full_text"] = text
            d = Document(page_content=text if text.strip() else source, metadata=m)
            docs.append(d)
        except Exception as e:
            print(f"OCR failed for {source}: {e}")
            # keep an entry so images can still be discovered by filename
            m = dict(meta_base)
            m["doc_type"] = "image"
            m["snippet"] = ""
            d = Document(page_content=source, metadata=m)
            docs.append(d)

    elif ext in [".mp4", ".mp3"]:
        try:
            print(f"Processing audio/video file: {source}")
            base, _ = os.path.splitext(saved_path)
            audio_path = base + "_audio.wav"
            if ext == ".mp4":
                video = VideoFileClip(saved_path)
                # write audio to wav
                video.audio.write_audiofile(audio_path, logger=None)
                video.close()
            elif ext == ".mp3":
                audio = AudioSegment.from_file(saved_path, format="mp3")
                audio.export(audio_path, format="wav")

            model = whisper.load_model("base")
            transcription = model.transcribe(audio_path, word_timestamps=False)  # keep segments
            # transcription contains 'segments' with start/end and 'text'
            segments = transcription.get("segments", [])
            if os.path.exists(audio_path):
                try:
                    os.remove(audio_path)
                except Exception:
                    pass

            # If segments exist, add one Document per segment
            if segments:
                for seg in segments:
                    seg_text = seg.get("text", "").strip()
                    if not seg_text:
                        continue
                    m = dict(meta_base)
                    m["doc_type"] = "video" if ext == ".mp4" else "audio"
                    m["start"] = float(seg.get("start", 0.0))
                    m["end"] = float(seg.get("end", 0.0))
                    m["transcript"] = seg_text
                    m["snippet"] = seg_text[:300]
                    d = Document(page_content=seg_text, metadata=m)
                    docs.append(d)
            else:
                # fallback: store full transcription as single doc
                text = transcription.get("text", "").strip()
                if text:
                    m = dict(meta_base)
                    m["doc_type"] = "video" if ext == ".mp4" else "audio"
                    m["snippet"] = text[:300]
                    m["transcript"] = text
                    d = Document(page_content=text, metadata=m)
                    docs.append(d)
        except Exception as e:
            print(f"Transcription failed for {source}: {e}")
            return []
    else:
        # Generic loader fallback: treat as text file
        try:
            with open(saved_path, "r", encoding="utf-8", errors="ignore") as fh:
                txt = fh.read()
                m = dict(meta_base)
                m["doc_type"] = "text"
                m["snippet"] = txt[:300]
                d = Document(page_content=txt, metadata=m)
                docs.append(d)
        except Exception as e:
            print(f"Generic load failed for {source}: {e}")
            return []

    # IMPORTANT: do NOT perform an unconditional split on docs (this is the root cause of CSV explosion).
    # Only split if you have a specific reason. We'll return the docs list as-is.
    return docs
